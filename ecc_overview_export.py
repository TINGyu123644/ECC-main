"""Export the full ECC overview (agents / skills / categories / three-mode usage) into a Word .docx file.

Run:  py -3 ecc_overview_export.py
Output: ECC-Overview.docx in the same directory.
"""

from __future__ import annotations

import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "ECC-Overview.docx")


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run_font(run, size_pt=10, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size_pt=14 if level == 1 else 12, bold=True, color=PRIMARY)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run, size_pt=10)
    return p


def add_table(doc, headers, rows, header_bg="1F4E79", col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        set_run_font(run, size_pt=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size_pt=9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    return table


# ---------------------------------------------------------------------------
# Content data
# ---------------------------------------------------------------------------

GENERAL_AGENTS = [
    ("planner", "实现规划（opus）"),
    ("architect", "系统架构与可扩展性（opus）"),
    ("code-architect", "现有代码库的特征架构蓝图"),
    ("code-explorer", "深度分析现有代码执行路径与依赖"),
    ("code-reviewer", "代码质量、可维护性、安全审查"),
    ("code-simplifier", "简化与重构（保留行为）"),
    ("tdd-guide", "测试驱动开发、TDD 工作流"),
    ("refactor-cleaner", "死代码清理（knip/depcheck/ts-prune）"),
    ("build-error-resolver", "构建/类型错误修复"),
    ("e2e-runner", "Playwright 端到端测试"),
    ("doc-updater", "更新 codemaps 与文档"),
    ("docs-lookup", "通过 Context7 查找文档"),
    ("performance-optimizer", "性能分析与优化"),
    ("silent-failure-hunter", "静默失败/吞错检查"),
    ("comment-analyzer", "注释质量与衰减风险"),
    ("pr-test-analyzer", "PR 测试覆盖率分析"),
    ("type-design-analyzer", "类型设计分析"),
    ("spec-miner", "棕地项目规格抽取（opus）"),
    ("agent-evaluator", "5 维度质量评分"),
    ("loop-operator", "自治循环运行与监控"),
    ("harness-optimizer", "harness 可靠性/成本/吞吐优化"),
    ("conversation-analyzer", "会话转录分析（为 hooks 找出可拦截行为）"),
    ("chief-of-staff", "多渠道通信分流（邮件/Slack 等）"),
]

LANG_PAIRS = [
    ("C/C++", "cpp-reviewer", "cpp-build-resolver"),
    ("C#/.NET", "csharp-reviewer", "—"),
    ("F#", "fsharp-reviewer", "—"),
    ("Dart/Flutter", "flutter-reviewer", "dart-build-resolver"),
    ("Go", "go-reviewer", "go-build-resolver"),
    ("Java（Spring/Quarkus）", "java-reviewer", "java-build-resolver"),
    ("Kotlin/Android/KMP", "kotlin-reviewer", "kotlin-build-resolver"),
    ("PHP/Laravel", "php-reviewer", "—"),
    ("Python/Django", "python-reviewer / django-reviewer", "django-build-resolver"),
    ("PyTorch", "—", "pytorch-build-resolver"),
    ("React", "react-reviewer", "react-build-resolver"),
    ("Rust", "rust-reviewer", "rust-build-resolver"),
    ("Swift", "swift-reviewer", "swift-build-resolver"),
    ("TypeScript/JS", "typescript-reviewer", "—"),
    ("Vue", "vue-reviewer", "—"),
    ("FastAPI", "fastapi-reviewer", "—"),
    ("HarmonyOS", "harmonyos-app-resolver", "—"),
]

SPECIALIST_AGENTS = [
    ("database-reviewer", "PostgreSQL/Supabase 数据库专家"),
    ("security-reviewer", "安全漏洞检测与修复"),
    ("seo-specialist", "SEO 技术审计与优化"),
    ("mle-reviewer", "ML 工程 / 模型上线 / 监控 / 回滚"),
    ("healthcare-reviewer", "医疗应用临床安全 / PHI 合规（opus）"),
    ("network-architect", "企业/多站点网络架构"),
    ("network-config-reviewer", "路由器/交换机配置审查"),
    ("network-troubleshooter", "网络连通性/DNS/路由诊断"),
    ("homelab-architect", "家庭/小型实验室网络规划"),
    ("a11y-architect", "WCAG 2.2 无障碍架构"),
    ("marketing-agent", "营销策略与文案"),
    ("opensource-forker", "开源分叉（去敏感信息）"),
    ("opensource-sanitizer", "开源发布前敏感信息扫描"),
    ("opensource-packager", "生成完整开源发布包"),
    ("gan-planner", "GAN Harness — 规划"),
    ("gan-generator", "GAN Harness — 生成实现"),
    ("gan-evaluator", "GAN Harness — 评估反馈"),
]

AGENT_SKILL_BINDINGS = [
    ("react-reviewer", "react-patterns, react-testing, accessibility"),
    ("react-build-resolver", "react-patterns, frontend-patterns"),
    ("vue-reviewer", "vue-patterns"),
    ("python-reviewer", "python-patterns, python-testing"),
    ("cpp-reviewer", "cpp-coding-standards, cpp-testing"),
    ("django-reviewer", "django-patterns, django-security, django-tdd, django-verification"),
    ("fastapi-reviewer", "fastapi-patterns"),
    ("springboot-reviewer", "springboot-patterns, springboot-security, springboot-tdd, springboot-verification"),
    ("laravel-reviewer", "laravel-patterns, laravel-security, laravel-tdd, laravel-verification"),
    ("rust-reviewer", "rust-patterns, rust-testing"),
    ("database-reviewer", "postgres-patterns, mysql-patterns, prisma-patterns"),
    ("mle-reviewer", "mle-workflow, eval-harness, benchmark-methodology"),
    ("healthcare-reviewer", "healthcare-emr-patterns, healthcare-cdss-patterns, healthcare-phi-compliance, hipaa-compliance, healthcare-eval-harness"),
    ("security-reviewer", "security-review, security-scan, security-bounty-hunter"),
    ("network-architect", "homelab-network-setup, homelab-network-readiness, network-bgp-diagnostics, network-config-validation"),
    ("a11y-architect", "accessibility, frontend-a11y"),
    ("tdd-guide", "tdd-workflow, cpp-testing/python-testing/rust-testing/golang-testing/..."),
    ("planner", "plan-canvas, plan-orchestrate, product-lens, product-capability"),
    ("e2e-runner", "e2e-testing, browser-qa, windows-desktop-e2e"),
    ("code-reviewer", "coding-standards, security-review"),
    ("marketing-agent", "marketing-campaign, brand-voice, seo"),
    ("opensource-forker/sanitizer/packager", "opensource-pipeline"),
    ("spec-miner", "codebase-onboarding, intent-driven-development"),
    ("harness-optimizer", "agent-architecture-audit, cost-aware-llm-pipeline"),
    ("loop-operator", "autonomous-loops, continuous-agent-loop"),
    ("agent-evaluator", "agent-self-evaluation, agent-eval"),
    ("conversation-analyzer", "hookify-rules"),
]

SKILL_CATEGORIES = [
    (
        "1. 语言/框架栈（language-stacks）",
        75,
        "按语言或框架切分的模式、测试、TDD、安全、验证子集",
        "python-patterns, react-patterns, vue-patterns, django-patterns, fastapi-patterns, "
        "springboot-patterns, kotlin-patterns, rust-patterns, laravel-patterns, quarkus-patterns, "
        "swiftui-patterns, hexagonal-architecture, jpa-patterns, prisma-patterns, redis-patterns, "
        "mysql-patterns, postgres-patterns 等",
    ),
    (
        "2. 研究与金融（research-finance）",
        21,
        "研究、检索、量化/链上/预测市场、科学数据库",
        "deep-research, iterative-retrieval, search-first, exa-search, recsys-pipeline-architect, "
        "scientific-db-pubmed-database, scientific-db-uspto-database, ito-trade-planner, "
        "prediction-market-oracle-research, defi-amm-security, llm-trading-agent-security",
    ),
    (
        "3. 运维与工具链（ops-tooling）",
        19,
        "GitHub/Jira、终端、网络、贸易合规、财务",
        "github-ops, jira-integration, terminal-ops, network-config-validation, security-scan, "
        "security-bounty-hunter, hookify-rules, gateguard, finance-billing-ops, "
        "logistics-exception-management, customs-trade-compliance",
    ),
    (
        "4. 营销与运营（marketing-ops）",
        17,
        "品牌、内容、SEO、邮件、社媒、投资人沟通",
        "marketing-campaign, brand-voice, seo, content-engine, social-publisher, email-ops, "
        "investor-outreach, mailtrap-email-integration, messages-ops",
    ),
    (
        "5. 元层/Harness（meta-harness）",
        16,
        "关于 Claude/LLM 本身的工程实践",
        "claude-devfleet, cost-aware-llm-pipeline, cost-tracking, context-budget, "
        "token-budget-advisor, council, verification-loop, benchmark-methodology, "
        "ai-regression-testing, ai-first-engineering, strategic-compact",
    ),
    (
        "6. 设计与多媒体（design-media）",
        16,
        "UI/UX、动效、视频、设计系统",
        "design-system, motion-foundations, motion-advanced, motion-patterns, motion-ui, "
        "manim-video, remotion-video-creation, video-editing, videodb, fal-ai-media, "
        "liquid-glass-design, ui-demo, taste",
    ),
    (
        "7. 智能体工程（agent-engineering）",
        16,
        "多 agent 编排、autonomy、continuous learning、self-eval",
        "agent-harness-construction, agent-architecture-audit, agent-self-evaluation, "
        "agentic-engineering, agentic-os, autonomous-agent-harness, continuous-learning, "
        "continuous-agent-loop, team-builder, team-agent-orchestration, loop-design-check",
    ),
    (
        "8. 工作流元数据（workflow-meta）",
        10,
        "项目/技能/会话/文档管理",
        "codebase-onboarding, code-tour, git-workflow, config-gc, skill-scout, skill-stocktake, "
        "skill-comply, product-capability, product-lens, project-flow-ops",
    ),
    (
        "9. 编排工作流（orchestration）",
        10,
        "PRP/orchestrator/plan-canvas/dmux",
        "orch-add-feature, orch-build-mvp, orch-change-feature, orch-fix-defect, orch-pipeline, "
        "orch-refine-code, plan-canvas, plan-orchestrate, ralphinho-rfc-pipeline, dmux-workflows",
    ),
    (
        "10. 网络运维（network-ops）",
        10,
        "家庭网络与运营商级网络",
        "homelab-network-setup, homelab-network-readiness, homelab-vlan-segmentation, "
        "homelab-wireguard-vpn, homelab-pihole-dns, network-bgp-diagnostics, "
        "network-interface-health, netmiko-ssh-automation, cisco-ios-patterns",
    ),
    (
        "11. 基础设施/数据（infra-data）",
        10,
        "数据库、容器、生产审计、部署",
        "kubernetes-patterns, docker-patterns, deployment-patterns, flox-environments, uncloud, "
        "bun-runtime, clickhouse-io, production-audit, terminal-ops, knowledge-ops",
    ),
    (
        "12. 医疗合规（healthcare）",
        5,
        "EMR/EHR、CDSS、PHI、HIPAA",
        "healthcare-emr-patterns, healthcare-cdss-patterns, healthcare-eval-harness, "
        "healthcare-phi-compliance, hipaa-compliance",
    ),
]

# Three-mode usage scenarios
SCENARIO_SKILL_ONLY = [
    ("「LLM 交易 bot 应该有哪些安全防护」", "llm-trading-agent-security", "知识问答，让父 Claude 读 SKILL.md 直接讲"),
    ("「TDD 工作流怎么走？覆盖率门槛多少」", "tdd-workflow", "解释方法论，无需执行"),
    ("「React 19 的 Server Component 怎么用」", "react-patterns", "概念讲解"),
    ("「Django 项目的 ORM 反模式有哪些」", "django-patterns", "清单类问题"),
    ("「医疗 PHI 数据有哪些合规要求」", "healthcare-phi-compliance + hipaa-compliance", "合规咨询"),
    ("「设计一个运动 App 的动效系统」", "motion-foundations + motion-patterns", "设计参考"),
    ("「我该选 PostgreSQL 还是 ClickHouse」", "postgres-patterns + clickhouse-io", "技术选型对比"),
    ("「如何审计一个交易智能合约」", "defi-amm-security", "审计方法论讲解"),
    ("「部署到 Kubernetes 怎么做」", "kubernetes-patterns + deployment-patterns", "部署流程说明"),
    ("「公司要发一篇技术博客」", "article-writing + content-engine", "文案方法论"),
]

SCENARIO_AGENT_ONLY = [
    ("「帮我把这段 Rust 代码简化一下」", "code-simplifier", "重构方法论内嵌在 prompt"),
    ("「构建报错了，帮我修」", "build-error-resolver", "错误修复流程完整"),
    ("「把项目里的死代码清掉」", "refactor-cleaner", "knip/depcheck/ts-prune 流程"),
    ("「写个 Playwright 端到端测试」", "e2e-runner", "测试生成流程"),
    ("「帮我规划这个功能的实现」", "planner", "拆解任务、识别依赖、排序"),
    ("「把整个项目架构画一下」", "architect", "系统设计方法"),
    ("「分析下这段对话，看哪些行为该被 hook 拦截」", "conversation-analyzer", "转录分析流程"),
    ("「优化这段代码的性能」", "performance-optimizer", "性能分析流程"),
    ("「从老代码里提取规格说明」", "spec-miner", "规格抽取流程（opus）"),
    ("「我要做开源，把敏感信息清掉」", "opensource-sanitizer", "20+ 模式扫描流程"),
]

SCENARIO_COMBO = [
    ("「审查这段 React 代码」", "react-reviewer", "react-patterns, react-testing, accessibility", "reviewer 知道审查流程，但需要 React 模式知识"),
    ("「审查这段 Vue 代码」", "vue-reviewer", "vue-patterns", "同上"),
    ("「审查这段 Python 代码」", "python-reviewer", "python-patterns, python-testing", "reviewer 需要 Pythonic 知识"),
    ("「审查 C++ 代码」", "cpp-reviewer", "cpp-coding-standards, cpp-testing", "编码规范要查"),
    ("「审查 Django 项目」", "django-reviewer", "django-patterns, django-security, django-tdd, django-verification", "4 个 skill 拼出 Django 完整审查"),
    ("「审查 Laravel 项目」", "laravel-reviewer", "laravel-patterns, laravel-security, laravel-tdd", "同上"),
    ("「审查 Spring Boot 项目」", "springboot-reviewer", "springboot-patterns, springboot-security, springboot-tdd, springboot-verification", "同上"),
    ("「审查 ML 流水线」", "mle-reviewer", "mle-workflow, eval-harness, benchmark-methodology", "ML 工程需要评测方法"),
    ("「审查医疗应用」", "healthcare-reviewer（opus）", "healthcare-emr-patterns, healthcare-cdss-patterns, healthcare-phi-compliance, hipaa-compliance, healthcare-eval-harness", "5 个 skill 拼临床安全全图"),
    ("「审查企业网络架构」", "network-architect", "homelab-network-setup, network-bgp-diagnostics, network-config-validation", "网络设计知识"),
    ("「无障碍设计审查」", "a11y-architect", "accessibility, frontend-a11y", "WCAG 标准"),
    ("「给 agent 打分」", "agent-evaluator", "agent-self-evaluation, agent-eval", "评测维度表"),
    ("「审查一段对话找 hook 候选」", "conversation-analyzer", "hookify-rules", "hook 规则知识"),
    ("「自治循环卡住了」", "loop-operator", "autonomous-loops, continuous-agent-loop", "loop 设计模式"),
    ("「优化 harness 配置」", "harness-optimizer", "agent-architecture-audit, cost-aware-llm-pipeline", "成本与可靠性知识"),
]


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    # Title
    title = doc.add_heading("ECC 项目结构总览", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, size_pt=22, bold=True, color=PRIMARY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Everything Claude Code  ·  Agents / Skills / Categories / Usage Patterns")
    set_run_font(sr, size_pt=11, italic=True, color=GREY)

    add_paragraph(
        doc,
        "项目路径：ECC-main/ECC-main（Claude Code 插件仓库，v2.0.0）。整体规模：67 个智能体 + 278 个技能 "
        "+ 94 个 slash 命令 + 多套 hooks/rules。",
    )

    # ===== Section 1: Agents =====
    add_heading(doc, "一、智能体（67 个，按功能分簇）", level=1)
    add_paragraph(
        doc,
        "智能体通过 agents/*.md 文件定义，frontmatter 中包含 name / description / tools / model 四个字段。"
        "模型分布：57 个 sonnet、4 个 opus、6 个 haiku。",
    )

    add_heading(doc, "1.1 通用与元层级（核心调度）", level=2)
    add_table(doc, ["智能体", "用途"], GENERAL_AGENTS, col_widths=[5.5, 11])

    add_heading(doc, "1.2 语言/栈专属审查与构建修复（pair 形式）", level=2)
    add_table(doc, ["语言/框架", "Reviewer", "Build Resolver"], LANG_PAIRS, col_widths=[5, 6, 6])

    add_heading(doc, "1.3 领域专家（specialist）", level=2)
    add_table(doc, ["智能体", "专长"], SPECIALIST_AGENTS, col_widths=[6, 11])

    # ===== Section 2: Agent ↔ Skill bindings =====
    add_heading(doc, "二、智能体 ↔ 技能调用关系", level=1)
    add_paragraph(
        doc,
        "关键事实：智能体并不通过 frontmatter 硬绑定技能。每个 agent 在 frontmatter 中声明的是工具白名单（tools），"
        "而技能（skill）是独立的、按需加载的工作流文档，由 agent 的 prompt 文本中以自然语言引用，并由主 Claude "
        "实例动态注入上下文。",
    )
    add_paragraph(doc, "调用机制：", bold=True)
    add_bullet(doc, "Agent 文件通常在结尾的 'Skills:' 一节写明推荐加载的技能路径。")
    add_bullet(doc, "当父 Claude 决定委派某个 agent 时，会按对应 skill 的目录（skills/<skill-name>/SKILL.md）读取并把内容塞入 agent 的上下文。")
    add_bullet(doc, "Agent 内部可用的工具集合（tools 字段）决定了它实际能执行什么，与技能加载是两条独立维度。")

    add_heading(doc, "2.1 典型绑定示例", level=2)
    add_table(doc, ["Agent", "推荐调用技能"], AGENT_SKILL_BINDINGS, col_widths=[6, 11])

    add_paragraph(
        doc,
        "其它未在 agent 中显式列出的技能被分类归到 commands、slash 入口或工作流编排层，"
        "由 commands/*.md 兼容垫片、workflows/、hooks/ 调用。",
        italic=True,
    )

    # ===== Section 3: Skill categories =====
    add_heading(doc, "三、技能分类（278 个，按命名约定与功能域）", level=1)
    add_paragraph(
        doc,
        "技能定义在 skills/<name>/SKILL.md，frontmatter 含 name / description / metadata.origin。"
        "官方「工作流面」政策：新贡献默认先放进 skills/，只有需要历史兼容时才在 commands/ 加 shim。",
    )

    rows = [(name, str(count), desc, examples) for name, count, desc, examples in SKILL_CATEGORIES]
    add_table(
        doc,
        ["类别", "数量", "说明", "典型技能"],
        rows,
        col_widths=[5, 1.6, 5.5, 7.5],
    )

    add_paragraph(doc, "其余 53 个跨类别的「通用 / 杂项」技能分散在以下主题：", bold=True)
    add_bullet(doc, "API 设计、TDD 工作流、错误处理、E2E 测试、代码库健康")
    add_bullet(doc, "MCP/集成、可观测性、可移植性（visa-doc-translate、windows-desktop-e2e）")
    add_bullet(doc, "开源流水线（opensource-pipeline）")
    add_bullet(doc, "ECC 自描述（ecc-guide、ecc-recipes、configure-ecc）")

    add_paragraph(doc, "技能命名约定：", bold=True)
    add_bullet(doc, "-patterns — 设计模式/最佳实践")
    add_bullet(doc, "-testing / -tdd — 测试方法学")
    add_bullet(doc, "-verification — 验证流程")
    add_bullet(doc, "-security — 安全审计")
    add_bullet(doc, "-coding-standards — 编码规范")
    add_bullet(doc, "-build / -build-fix — 构建修复")
    add_bullet(doc, "orch-* — 多 agent 编排工作流")
    add_bullet(doc, "homelab-* / network-* — 网络基础设施")
    add_bullet(doc, "healthcare-* / hipaa-* — 医疗合规")

    # ===== Section 4: Three-mode usage patterns =====
    add_heading(doc, "四、三种使用模式（Skill 与 Agent 的实际组合方式）", level=1)
    add_paragraph(
        doc,
        "在实际运行中，Skill 与 Agent 的组合并非固定不变。父 Claude 在每次对话中根据用户提问的性质，"
        "在以下三种模式中独立决策：",
    )

    # --- Mode 1: Skill only ---
    add_heading(doc, "4.1 模式 A：只调 Skill（父 Claude 自答）", level=2)
    add_paragraph(
        doc,
        "特征：用户要的是知识 / 方法 / 清单，不是「动手做」。父 Claude 读取 SKILL.md 后直接回答，不委派 agent。",
        italic=True,
    )
    add_table(
        doc,
        ["用户提问", "加载的 Skill", "为什么不派 agent"],
        SCENARIO_SKILL_ONLY,
        col_widths=[7, 5.5, 6],
    )

    # --- Mode 2: Agent only ---
    add_heading(doc, "4.2 模式 B：只调 Agent（Agent 自带流程）", level=2)
    add_paragraph(
        doc,
        "特征：任务可执行、可验证、有明确产物，agent 的内置流程够用。",
        italic=True,
    )
    add_table(
        doc,
        ["用户提问", "派出的 Agent", "Agent 自带的能力"],
        SCENARIO_AGENT_ONLY,
        col_widths=[7, 5.5, 6],
    )

    # --- Mode 3: Combo ---
    add_heading(doc, "4.3 模式 C：Agent + Skill 组合（Skill 增强 Agent）", level=2)
    add_paragraph(
        doc,
        "特征：任务既需要专业角色又需要深度领域知识。Agent 知道「做什么」，Skill 告诉它「具体怎么做对」。",
        italic=True,
    )
    add_table(
        doc,
        ["用户提问", "Agent", "注入的 Skill", "为什么要组合"],
        SCENARIO_COMBO,
        col_widths=[5, 4, 5, 5],
    )

    # --- Decision mnemonic ---
    add_heading(doc, "4.4 三种模式的判定口诀", level=2)
    add_paragraph(
        doc,
        "用户提问 → 「是什么 / 为什么 / 怎么做（讲方法）」→ 只调 skill；"
        "「帮我做 / 帮我修 / 帮我跑（动手）」→ 只调 agent（除非任务需要深度领域知识）；"
        "「审查 X 栈的代码 / 评估 X 领域的方案 / 设计 X 系统」→ agent + skill 组合（这是审查 / 评估类任务的标配）。",
    )
    add_paragraph(doc, "三类比例（经验值）：", bold=True)
    add_bullet(doc, "只调 skill：~40%（父 Claude 自答类，多见于知识问答、方法论咨询）")
    add_bullet(doc, "只调 agent：~35%（重构、构建修复、端到端测试等执行型任务）")
    add_bullet(doc, "Agent + Skill 组合：~25%（审查类 agent 居多，因为审查需要既懂流程又有领域知识）")

    # --- Key insight ---
    add_heading(doc, "4.5 一个值得注意的细节", level=2)
    add_paragraph(
        doc,
        "所有「代码审查」类 agent 几乎都是组合模式（模式 C）—— 这是 ECC 的设计哲学：",
        italic=True,
    )
    add_bullet(doc, "通用 code-reviewer 用 coding-standards + security-review")
    add_bullet(doc, "各语言 reviewer（python / vue / cpp / django / ...）都用对应 *-patterns + *-testing + *-security")
    add_bullet(doc, "领域 reviewer（healthcare / mle / network / ...）都用对应的 *-patterns + 合规 / 方法论 skill")
    add_paragraph(
        doc,
        "审查这件事 = 流程（agent 负责）+ 领域知识（skill 负责），缺一不可，所以必然是组合模式。",
        bold=True,
    )

    # ===== Section 5: Additional info =====
    add_heading(doc, "五、补充信息", level=1)
    add_bullet(doc, "Slash 命令（commands/）：94 个，历史上是技能入口，现在保留为兼容垫片（shim），长期方向是 skills-first。")
    add_bullet(doc, "Hooks（hooks/ + scripts/hooks/）：触发式自动化，ECC 内置了 gateguard、hookify、fact-forcing gate 等。")
    add_bullet(doc, "Rules（rules/）：始终遵循的编码/安全准则，按语言拆分。")
    add_bullet(doc, "MCP 配置（mcp-configs/）：14 个外部服务集成（Context7、Playwright 等）。")
    add_bullet(doc, "ECC 自描述入口：技能 ecc-guide 会在用户询问 ECC 自身结构时读取实时仓库并给出答案。")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote: {out}")
    print(f"Size: {os.path.getsize(out):,} bytes")