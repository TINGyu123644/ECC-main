"""Export the ECC skill/agent relationship discussion into a Word .docx file.

Run:  py -3 ecc_relationship_export.py
Output: ECC-Skill-Agent-Relationship.docx in the same directory.
"""

from __future__ import annotations

import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "ECC-Skill-Agent-Relationship.docx")


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
SOFT_BG = "EAF2FB"
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run_font(run, size_pt=10, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "Microsoft YaHei"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size_pt=14 if level == 1 else 12, bold=True, color=color or PRIMARY)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run, size_pt=10)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_run_font(run, size_pt=10)
    return p


def add_callout(doc, text, bg="EAF2FB"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5, italic=True, color=PRIMARY)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, header_bg="1F4E79", col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        set_run_font(run, size_pt=10, bold=True, color=WHITE)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size_pt=9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    return table


# ---------------------------------------------------------------------------
# Content data
# ---------------------------------------------------------------------------

CORE_DIFF_TABLE = [
    ("本质", "一份知识/工作流文档", "一个带工具的角色定义"),
    ("有工具吗", "❌ 没有", "✅ 有 tools 白名单（Read/Grep/Bash…）"),
    ("能跑代码吗", "❌ 自己不能", "✅ 能自己调工具"),
    ("frontmatter", "name + description + metadata",
                   "name + description + tools + model"),
    ("存储位置", "skills/<name>/SKILL.md", "agents/<name>.md"),
    ("数量", "278", "67"),
    ("类比", "一本手册", "一个工种的岗位说明书"),
]

SKILL_ONLY_SCENARIOS = [
    ("LLM 交易 bot 应该有哪些安全防护", "llm-trading-agent-security", "知识问答"),
    ("TDD 工作流怎么走？覆盖率门槛多少", "tdd-workflow", "解释方法论"),
    ("React 19 的 Server Component 怎么用", "react-patterns", "概念讲解"),
    ("Django 项目的 ORM 反模式有哪些", "django-patterns", "清单类问题"),
    ("医疗 PHI 数据有哪些合规要求", "healthcare-phi-compliance + hipaa-compliance", "合规咨询"),
    ("设计一个运动 App 的动效系统", "motion-foundations + motion-patterns", "设计参考"),
    ("我该选 PostgreSQL 还是 ClickHouse", "postgres-patterns + clickhouse-io", "技术选型对比"),
    ("如何审计一个交易智能合约", "defi-amm-security", "审计方法论"),
    ("部署到 Kubernetes 怎么做", "kubernetes-patterns + deployment-patterns", "部署流程"),
    ("公司要发一篇技术博客", "article-writing + content-engine", "文案方法论"),
]

AGENT_ONLY_SCENARIOS = [
    ("帮我把这段 Rust 代码简化一下", "code-simplifier", "重构方法论内嵌"),
    ("构建报错了，帮我修", "build-error-resolver", "错误修复流程"),
    ("把项目里的死代码清掉", "refactor-cleaner", "knip/depcheck 流程"),
    ("写个 Playwright 端到端测试", "e2e-runner", "测试生成流程"),
    ("帮我规划这个功能的实现", "planner", "拆解任务、识别依赖"),
    ("把整个项目架构画一下", "architect", "系统设计方法"),
    ("分析这段对话，看哪些行为该被 hook 拦截", "conversation-analyzer", "转录分析流程"),
    ("优化这段代码的性能", "performance-optimizer", "性能分析流程"),
    ("从老代码里提取规格说明", "spec-miner", "规格抽取（opus）"),
    ("我要做开源，把敏感信息清掉", "opensource-sanitizer", "20+ 模式扫描"),
]

COMBINED_SCENARIOS = [
    ("审查这段 React 代码", "react-reviewer",
     "react-patterns, react-testing, accessibility",
     "reviewer 知道审查流程，但需要 React 模式知识"),
    ("审查这段 Vue 代码", "vue-reviewer", "vue-patterns", "同上"),
    ("审查这段 Python 代码", "python-reviewer", "python-patterns, python-testing", "需要 Pythonic 知识"),
    ("审查 C++ 代码", "cpp-reviewer", "cpp-coding-standards, cpp-testing", "编码规范要查"),
    ("审查 Django 项目", "django-reviewer",
     "django-patterns, django-security, django-tdd, django-verification",
     "4 个 skill 拼出 Django 完整审查"),
    ("审查 Laravel 项目", "laravel-reviewer",
     "laravel-patterns, laravel-security, laravel-tdd",
     "同上"),
    ("审查 Spring Boot 项目", "springboot-reviewer",
     "springboot-patterns, springboot-security, springboot-tdd, springboot-verification",
     "同上"),
    ("审查 ML 流水线", "mle-reviewer", "mle-workflow, eval-harness, benchmark-methodology",
     "ML 工程需要评测方法"),
    ("审查医疗应用", "healthcare-reviewer (opus)",
     "healthcare-emr-patterns, healthcare-cdss-patterns, healthcare-phi-compliance, hipaa-compliance, healthcare-eval-harness",
     "5 个 skill 拼临床安全全图"),
    ("审查企业网络架构", "network-architect",
     "homelab-network-setup, network-bgp-diagnostics, network-config-validation",
     "网络设计知识"),
    ("无障碍设计审查", "a11y-architect", "accessibility, frontend-a11y", "WCAG 标准"),
    ("给 agent 打分", "agent-evaluator", "agent-self-evaluation, agent-eval", "评测维度表"),
    ("审查一段对话找 hook 候选", "conversation-analyzer", "hookify-rules", "hook 规则知识"),
    ("自治循环卡住了", "loop-operator", "autonomous-loops, continuous-agent-loop", "loop 设计模式"),
    ("优化 harness 配置", "harness-optimizer",
     "agent-architecture-audit, cost-aware-llm-pipeline", "成本与可靠性知识"),
]


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    title = doc.add_heading("Skill 与 Agent 的关系", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, size_pt=22, bold=True, color=PRIMARY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Everything Claude Code  ·  三种调度模式与场景对照")
    set_run_font(sr, size_pt=11, italic=True, color=GREY)

    # Section 1
    add_heading(doc, "一、Skill vs Agent：本职差异", level=1)
    add_paragraph(
        doc,
        "一句话：Skill 是知识手册，Agent 是带工具的执行角色。两者都是 markdown 文档，"
        "父 Claude 是唯一把两者连起来的调度者。",
    )
    add_table(doc, ["维度", "Skill", "Agent"], CORE_DIFF_TABLE,
              col_widths=[3.5, 6.5, 7])

    # Section 2
    add_heading(doc, "二、三种调度模式并存", level=1)
    add_paragraph(
        doc,
        "用户每一次提问，父 Claude 都会做一次判断，落到下面三种模式之一：",
    )
    add_bullet(doc, "只调 Skill（父 Claude 自答）：知识/方法类问题")
    add_bullet(doc, "只调 Agent（执行类任务）：动手做/修/跑/测")
    add_bullet(doc, "Agent + Skill 组合（增强型）：审查/评估/设计类深度任务")

    add_callout(
        doc,
        "核心结论：Skill 是可选的增强材料，Agent 是可选的执行载体，"
        "父 Claude 决定是否使用以及如何组合。三种模式都常见，没有谁占绝对主流。",
    )

    # Section 3
    add_heading(doc, "三、场景对照表（10+10+15 个真实例子）", level=1)

    add_heading(doc, "3.1 只调 Skill（父 Claude 自答）", level=2)
    add_paragraph(doc, "触发条件：用户要的是知识/方法/清单，不是动手做。")
    add_table(
        doc,
        ["用户提问", "加载的 Skill", "为什么不派 agent"],
        SKILL_ONLY_SCENARIOS,
        col_widths=[5.5, 6.5, 5],
    )

    add_heading(doc, "3.2 只调 Agent（Agent 自带流程）", level=2)
    add_paragraph(doc, "触发条件：任务可执行、可验证、有明确产物，agent 的内置流程够用。")
    add_table(
        doc,
        ["用户提问", "派出的 Agent", "Agent 自带的能力"],
        AGENT_ONLY_SCENARIOS,
        col_widths=[5.5, 5, 6.5],
    )

    add_heading(doc, "3.3 Agent + Skill 组合（Skill 增强 Agent）", level=2)
    add_paragraph(
        doc,
        "触发条件：任务既需要专业角色又需要深度领域知识。"
        "agent 知道「做什么」，skill 告诉它「具体怎么做对」。",
    )
    add_table(
        doc,
        ["用户提问", "Agent", "注入的 Skill", "为什么要组合"],
        COMBINED_SCENARIOS,
        col_widths=[4.5, 3.5, 5, 4],
    )

    # Section 4
    add_heading(doc, "四、判定口诀（父 Claude 视角）", level=1)
    add_paragraph(doc, "用户问什么？", bold=True)
    add_numbered(doc, '是 "是什么/为什么/怎么做（讲方法）" → 只调 skill')
    add_numbered(doc, '是 "帮我做/帮我修/帮我跑（动手）" → 只调 agent（除非任务需要深度领域知识）')
    add_numbered(doc, '是 "审查 X 栈的代码 / 评估 X 领域的方案 / 设计 X 系统" → agent + skill 组合')

    add_heading(doc, "大致占比", level=2)
    add_table(
        doc,
        ["模式", "占比", "典型任务类型"],
        [
            ("只调 Skill", "~40%", "知识问答、方法论咨询、技术选型、合规咨询"),
            ("只调 Agent", "~35%", "重构、构建修复、端到端测试、执行类操作"),
            ("Agent + Skill 组合", "~25%", "代码审查、领域评估、架构设计"),
        ],
        col_widths=[4.5, 2, 10.5],
    )

    # Section 5
    add_heading(doc, "五、为什么审查类 Agent 几乎都是组合模式？", level=1)
    add_paragraph(
        doc,
        "所有「代码审查」类 agent 几乎都是 agent + skill 组合（类型 3），"
        "这是 ECC 的设计哲学：",
    )
    add_bullet(doc, "通用 code-reviewer 用 coding-standards + security-review")
    add_bullet(doc, "各语言 reviewer（python/vue/cpp/django/...）都用对应 *-patterns + *-testing + *-security")
    add_bullet(doc, "领域 reviewer（healthcare/mle/network/...）都用对应的 *-patterns + 合规/方法论 skill")
    add_paragraph(
        doc,
        "审查这件事 = 流程（agent 负责）+ 领域知识（skill 负责），缺一不可，"
        "所以必然是组合模式。",
    )

    # Section 6
    add_heading(doc, "六、常见误解澄清", level=1)
    add_paragraph(doc, "误解 1：「Skill 找到对应的 Agent，Agent 来做事」", bold=True)
    add_paragraph(
        doc,
        "❌ 错。Skill 和 Agent 没有任何注册表或路由表。父 Claude 才是把它们配对起来的唯一角色。"
        "Skill 的 description 里从不点名 agent，Agent 的 Skills: 段也只是软建议。",
    )
    add_paragraph(doc, "误解 2：「Agent 是 Skill 的管家，Agent 服务 Skill」", bold=True)
    add_paragraph(
        doc,
        "❌ 方向反了。Skill 是被引用的「手册」，Agent 是主动使用手册的「工匠」。"
        "更准确的比喻：Skill 是工具箱，Agent 是工匠；工匠挑工具干活，不是工具的管家。",
    )
    add_paragraph(doc, "误解 3：「Skill 和 Agent 完全独立，没有关系」", bold=True)
    add_paragraph(
        doc,
        "❌ 也不全对。大约 10% 的 agent 会在 prompt 末尾显式列出推荐 skill（react-reviewer、"
        "vue-reviewer、react-build-resolver、agent-evaluator、a11y-architect 等），"
        "这是「软约定」——不是强制，但反映了设计意图。",
    )

    add_callout(
        doc,
        "最准确的一句话：Skill 和 Agent 是两类正交资源，由父 Claude 在一次判断里同时识别、灵活组合。"
        "三者形成 Skill（手册）+ Agent（工匠）+ 父 Claude（项目经理）的三层架构。",
    )

    # Section 7
    add_heading(doc, "七、关系图（ASCII 示意）", level=1)
    diagram = (
        "            ┌──────────────┐\n"
        "            │   用户提问    │\n"
        "            └──────┬───────┘\n"
        "                   ↓\n"
        "            ┌──────────────┐\n"
        "            │   父 Claude   │   ← 唯一的「连接者」\n"
        "            │  (调度员)     │\n"
        "            └──────┬───────┘\n"
        "                   │ 运行时决策\n"
        "        ┌──────────┴──────────┐\n"
        "        ↓                     ↓\n"
        "   ┌─────────┐           ┌─────────┐\n"
        "   │ Skill A │           │ Agent 1 │   ← Agent 可在 prompt 里\n"
        "   │  (手册)  │           │  (工匠)  │      「建议」用哪些 skill\n"
        "   └─────────┘           └─────────┘\n"
        "        ↑                     │\n"
        "        │  内容注入（可选）    │\n"
        "        └─────────────────────┘\n"
        "        ↑\n"
        "        │ 父 Claude 决定注入什么\n"
        "   ┌─────────┐\n"
        "   │ 用户请求 │\n"
        "   └─────────┘\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)

    # Section 8
    add_heading(doc, "八、最终结论（一句话）", level=1)
    add_callout(
        doc,
        "Skill 是被动文档，Agent 是主动工人，父 Claude 是调度员。"
        "没有「skill 路由到 agent」这种事；要么父 Claude 自己读 skill，"
        "要么父 Claude 把 skill 喂给某个 agent，要么都不做（agent 独立工作）。",
        bg="FFF6E5",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote: {out}")
    print(f"Size: {os.path.getsize(out):,} bytes")